import pandas as pd
import xgboost as xgb
import numpy as np
import sklearn
from sklearn.preprocessing import OneHotEncoder
from sklearn.impute import SimpleImputer
from imblearn.over_sampling import SMOTE
from imblearn.pipeline import make_pipeline
from imblearn.combine import SMOTEENN
from imblearn.over_sampling import SMOTENC
import matplotlib.pyplot as plt
from PIL import Image
from scipy.interpolate import BSpline, make_interp_spline, interp1d
# import rpy2.robjects as robjects
# from rpy2.robjects.packages import importr
import csv
from dfply import *
from xgboost import XGBClassifier
import itertools
import os
from sklearn.model_selection import RandomizedSearchCV, GridSearchCV
from sklearn.metrics import roc_auc_score
from sklearn.model_selection import StratifiedKFold
import time
import pickle
import math
import matplotlib.pyplot as plt
from glob import glob

import importlib
import time
import requests

from joblib import Parallel, delayed
from joblib import parallel_backend
import seaborn as sns

from catboost import Pool, cv
import shap as sp
from os.path import exists
import itertools
from scipy.interpolate import BSpline, make_interp_spline, interp1d
from scipy import interpolate
import json
from matplotlib.lines import Line2D
import dataframe_image

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os
import sys
import inspect

currentdir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
parentdir = os.path.dirname(currentdir)
sys.path.insert(0, parentdir) 

import utils_code2name
importlib.reload(utils_code2name)

class plotmeta:
    def __init__(self, order='single', interaction = 'nointeraction', additional_parameter=None):
        self.datafolder = '/home/hoyinchan/blue/Data/data2022/'
        self.home_directory = "/home/hoyinchan/code/AKI_CDM_PY/MetaRegression/"
        self.home_data_directory = "/blue/yonghui.wu/hoyinchan/program_data/AKI_CDM_PY/MetaRegression/"
        self.order = order
        self.interaction = interaction
        self.filename = "gamalltmp_"+order+"_"+interaction
        if additional_parameter is not None:
            self.filename = self.filename+"_"+additional_parameter
        self.fits = None
        self.gamdata_fitdata = dict()
        self.gamdata_plotdata1d = dict()
        self.gamdata_plotdata2d = dict()
        self.gamdata_plotdatag = dict()
        self.gamdata_plotdatap = dict()
        self.shapdf = None
        self.plot_range = None
#        self.cattarget = ["PX:CH:J1940", "PX:09:96.72"]
        self.translator = utils_code2name.code2name()

        cattargetdf = pd.read_parquet('/home/hoyinchan/code/AKI_CDM_PY/bool_columns.parquet')
        self.cattarget = cattargetdf['index']
        self.translate_omop_pre()
        
    def load_raw_data(self):
        self.shapdf = pd.read_parquet('/home/hoyinchan/blue/Data/data2022/'+'shapalltmp.parquet')
        self.features = [x for x in np.unique(list(self.shapdf.columns.str.split('_').str[0])) if x != 'site']
        self.features = ['ORIGINAL_BMI' if feature == 'ORIGINAL' else feature for feature in self.features]        
            
    def get_meta_data(self, filename="gamalltmp_single.json"):
        if self.filename == 'gamalltmp_double_interaction':
            gamdata = pd.read_pickle('/home/hoyinchan/blue/program_data/AKI_CDM_PY/MetaRegression/gamalltmp_double_interaction_noAUC_json'+".pkl")
            gamdata.columns= [-1,0,1,2]
            gamdata = gamdata.explode(-1)
            self.get_meta_data_sub(gamdata, 'double_interaction')
        else:
#            gamdata = pd.read_json(self.filename+".json")
            gamdataX = pd.read_json(filename)
            for i in range(gamdataX.shape[1]):
                self.get_meta_data_sub(pd.DataFrame(list(gamdataX[i])), gamdataX[i][0][1][0])
    
    def get_meta_data_sub_old(self, gamdata2, label):
        gamdata = gamdata2.copy()
        if label != 'double_interaction':        
            gamdata = gamdata.drop([1],axis=1)
            gamdata.columns = [0,1,2,3]
                       
        gamdata = gamdata.explode(0)
        gamdata_fitdata = gamdata.copy().drop([2],axis=1)
        gamdata_fitdata = gamdata.copy().drop([3],axis=1, errors='ignore')        
        gamdata_fitdata = gamdata_fitdata.join(pd.json_normalize(gamdata_fitdata[1]))
        self.gamdata_fitdata[label] = gamdata_fitdata.drop([1],axis=1)            
        self.gamdata_fitdata[label].index = self.gamdata_fitdata[label][0]
        
        gamdata_plotdata = gamdata.copy().drop([1],axis=1)
        gamdata_fitdata = gamdata.copy().drop([3],axis=1, errors='ignore')        
        gamdata_plotdata = gamdata_plotdata.explode(2)
        gamdata_plotdata['keys'] = [list(x.keys()) for x in gamdata_plotdata[2]]          

        plot2d = np.logical_and(['x' in k for k in gamdata_plotdata['keys']], ['y' in k for k in gamdata_plotdata['keys']])
        plotg = np.logical_and(['x' not in k for k in gamdata_plotdata['keys']], ['y' not in k for k in gamdata_plotdata['keys']])
        plot1d = np.logical_and(['x' in k for k in gamdata_plotdata['keys']], ['y' not in k for k in gamdata_plotdata['keys']])
        
        gamdata_plotdata1 = gamdata_plotdata[plot1d].reset_index(drop=True)
        gamdata_plotdata2 = gamdata_plotdata[plotg].reset_index(drop=True)      
        gamdata_plotdata3 = gamdata_plotdata[plot2d].reset_index(drop=True)               

        if not gamdata_plotdata1.empty:
            gamdata_plotdata1 = gamdata_plotdata1.join(pd.json_normalize(gamdata_plotdata1[2]))
            gamdata_plotdata1['fit'] = [np.array(x).flatten() for x in gamdata_plotdata1['fit']]
            gamdata_plotdata1 = gamdata_plotdata1.explode('xlab')
            gamdata_plotdata1 = gamdata_plotdata1.explode('ylab')
            self.gamdata_plotdata1d[label] = gamdata_plotdata1

        if not gamdata_plotdata2.empty:
            gamdata_plotdata2 = gamdata_plotdata2.join(pd.json_normalize(gamdata_plotdata2[2]))
            gamdata_plotdata2['fit'] = [np.array(x).flatten() for x in gamdata_plotdata2['fit']]
            self.gamdata_plotdatag[label] = gamdata_plotdata2      

        if not gamdata_plotdata3.empty:
#            gamdata_plotdata3 = gamdata_plotdata3.join(pd.json_normalize(gamdata_plotdata3[2]))
            gamdata_plotdata3 = gamdata_plotdata3.reset_index(drop=True).join(pd.json_normalize(gamdata_plotdata3[2]))
            gamdata_plotdata3['fit'] = [np.array(x).flatten() for x in gamdata_plotdata3['fit']]
            gamdata_plotdata3 = gamdata_plotdata3.explode('xlab')
            gamdata_plotdata3 = gamdata_plotdata3.explode('ylab')
            self.gamdata_plotdata2d[label] = gamdata_plotdata3        

        try:
            gamdata_plotdatap = gamdata.copy().drop([1,2],axis=1)
            gamdata_plotdatap[3] = [pd.DataFrame(x['Name']) for x in gamdata_plotdatap[3]]
            self.gamdata_plotdatap[label] = gamdata_plotdatap
        except:
            pass

    def get_meta_data_sub(self, gamdata2, label):
        gamdata = gamdata2.copy()
        if label != 'double_interaction':        
            gamdata = gamdata.drop([1],axis=1)
            gamdata.columns = [0,1,2,3]
                       
        gamdata = gamdata.explode(0)
        gamdata_fitdata = gamdata.copy().drop([2],axis=1)
        gamdata_fitdata = gamdata.copy().drop([3],axis=1, errors='ignore')        
        gamdata_fitdata = gamdata_fitdata.join(pd.json_normalize(gamdata_fitdata[1]))
        self.gamdata_fitdata[label] = gamdata_fitdata.drop([1],axis=1)            
        self.gamdata_fitdata[label].index = self.gamdata_fitdata[label][0]
        
        gamdata_plotdata = gamdata.copy().drop([1],axis=1)
        gamdata_fitdata = gamdata.copy().drop([3],axis=1, errors='ignore')        
        gamdata_plotdata = gamdata_plotdata.explode(2)
        gamdata_plotdata['keys'] = [list(x.keys()) for x in gamdata_plotdata[2]]          

        plot2d = np.logical_and(['x' in k for k in gamdata_plotdata['keys']], ['y' in k for k in gamdata_plotdata['keys']])
        plotg = np.logical_and(['x' not in k for k in gamdata_plotdata['keys']], ['y' not in k for k in gamdata_plotdata['keys']])
        plot1d = np.logical_and(['x' in k for k in gamdata_plotdata['keys']], ['y' not in k for k in gamdata_plotdata['keys']])
        
        gamdata_plotdata1 = gamdata_plotdata[plot1d].reset_index(drop=True)
        gamdata_plotdata2 = gamdata_plotdata[plotg].reset_index(drop=True)      
        gamdata_plotdata3 = gamdata_plotdata[plot2d].reset_index(drop=True)               

        if not gamdata_plotdata1.empty:
            gamdata_plotdata1 = gamdata_plotdata1.join(pd.json_normalize(gamdata_plotdata1[2]))
            gamdata_plotdata1['fit'] = [np.array(x).flatten() for x in gamdata_plotdata1['fit']]
            gamdata_plotdata1 = gamdata_plotdata1.explode('xlab')
            gamdata_plotdata1 = gamdata_plotdata1.explode('ylab')
            self.gamdata_plotdata1d[label] = gamdata_plotdata1

        if not gamdata_plotdata2.empty:
            gamdata_plotdata2 = gamdata_plotdata2.join(pd.json_normalize(gamdata_plotdata2[2]))
            gamdata_plotdata2['fit'] = [np.array(x).flatten() for x in gamdata_plotdata2['fit']]
            self.gamdata_plotdatag[label] = gamdata_plotdata2      

        if not gamdata_plotdata3.empty:
#            gamdata_plotdata3 = gamdata_plotdata3.join(pd.json_normalize(gamdata_plotdata3[2]))
            gamdata_plotdata3 = gamdata_plotdata3.reset_index(drop=True).join(pd.json_normalize(gamdata_plotdata3[2]))
            gamdata_plotdata3['fit'] = [np.array(x).flatten() for x in gamdata_plotdata3['fit']]
            gamdata_plotdata3 = gamdata_plotdata3.explode('xlab')
            gamdata_plotdata3 = gamdata_plotdata3.explode('ylab')
            self.gamdata_plotdata2d[label] = gamdata_plotdata3        

        try:
            gamdata_plotdatap = gamdata.copy().drop([1,2],axis=1)
            gamdata_plotdatap[3] = [pd.DataFrame(x['Name']) for x in gamdata_plotdatap[3]]
            self.gamdata_plotdatap[label] = gamdata_plotdatap
        except:
            pass        
        
    def calculate_local_fit(self, feature):
        
        print(f"Fitting: {feature}")
        columns_select = [feature+'_Names', feature+'_vals', 'site_m', 'site_d']
        shapdf = self.shapdf[columns_select].dropna()
        shapdf.columns = ['Name', 'val', 'site_m', 'site_d']
        shapdf['Feature'] = feature
        shapdf = shapdf.sort_values('Name')
        # shapdf = self.shapdf[self.shapdf['Feature']==feature].dropna().sort_values('Name')
        # shapdf = shapdf.sort_values('Name')
        def interpolate_raw(df):
            return interpolate.UnivariateSpline(df['Name'], df['val'],k=5)
#        raw_fits = shapdf.groupby(['site_d', 'site_m', 'Feature']).apply(interpolate_raw).reset_index()
        raw_fits = shapdf.groupby(['site_d', 'site_m', 'Feature']).apply(interpolate_raw).reset_index()
        return [feature, raw_fits]

    def calculate_local_fit_all(self, n_jobs=4):
        #self.fits = Parallel(n_jobs=n_jobs)(delayed(self.calculate_local_fit)(feature) for feature in self.shapdf['Feature'].unique())
#        self.fits = [self.calculate_local_fit(feature) for feature in self.shapdf['Feature'].unique()]
        self.fits = [self.calculate_local_fit(feature) for feature in self.features]
        self.fits = pd.DataFrame(self.fits)
    
    def plot_meta1(self, feature, filter_sitem=['MCRI'], ext_obj1=None, ext_obj2=None, ax=None, alpha=0.1, axis_label=True, margin=1):
#        raw_fits = self.calculate_local_fit(feature)
        gfc = self.gen_fit_character()
        raw_fits = self.fits[self.fits[0]==feature].iloc[0,1]
        for ffea in filter_sitem:
            raw_fits = raw_fits[raw_fits['site_m']!=ffea]
#            raw_fits = raw_fits[raw_fits['site_m']!=filter_sitem]
#        plt.figure(figsize=(16,16), dpi=400) 
        if ax is None:
            ax = plt.gca() 
        ax.axhline(y=0, linestyle='--', color='r')

        #Note: The gamdata_plotdataX.loc[:,'y'] fit is generated without the intercept, the plot has to add back the intercept for correct value
        
        gamdata_plotdataX = self.gamdata_plotdata1d['spline'][self.gamdata_plotdata1d['spline'][0]==feature].reset_index()
        spline_x = np.array(gamdata_plotdataX.loc[:,'x'].iloc[0])
        

        ax.plot(spline_x, spline_x*(gfc[gfc[0]==feature]['slope_linear'].iloc[0])+gfc[gfc[0]==feature]['intercept_linear'].iloc[0], linewidth=5, color='y', label='linear')
#        gamdata_plotdataX = self.gamdata_plotdatap['linear'][self.gamdata_plotdatap['linear'][0]==feature].reset_index().loc[0,3]        
#        ax.plot(gamdata_plotdataX.loc[:,'x'], gamdata_plotdataX.loc[:,'y']+gfc[gfc[0]==feature]['intercept_linear'].iloc[0], linewidth=5, color='y', label='linear')
#        ax.plot(gamdata_plotdataX.loc[:,'x'], gamdata_plotdataX.loc[:,'x']*gfc[gfc[0]==feature]['slope_linear'].iloc[0]+
#                                                gfc[gfc[0]==feature]['intercept_linear'].iloc[0], linewidth=5, color='y', label='linear')


        ax.plot(spline_x, spline_x**2*(gfc[gfc[0]==feature]['curvature_quadratic'].iloc[0])+spline_x*(gfc[gfc[0]==feature]['slope_quadratic'].iloc[0])+gfc[gfc[0]==feature]['intercept_quadratic'].iloc[0], linewidth=5, color='g', label='quadratic')
#        gamdata_plotdataX = self.gamdata_plotdatap['quadratic'][self.gamdata_plotdatap['quadratic'][0]==feature].reset_index().loc[0,3]        
#        ax.plot(gamdata_plotdataX.loc[:,'x'], gamdata_plotdataX.loc[:,'y']+gfc[gfc[0]==feature]['intercept_quadratic'].iloc[0], linewidth=5, color='g', label='quadratic')
        # ax.plot(gamdata_plotdataX.loc[:,'x'], np.square(gamdata_plotdataX.loc[:,'x'])*gfc[gfc[0]==feature]['curvature_quadratic'].iloc[0]+
        #                                         gamdata_plotdataX.loc[:,'x']*gfc[gfc[0]==feature]['slope_quadratic'].iloc[0]+
        #                                         gfc[gfc[0]==feature]['intercept_quadratic'].iloc[0], linewidth=5, color='g', label='quadratic')
        
        gamdata_plotdataX = self.gamdata_plotdata1d['spline'][self.gamdata_plotdata1d['spline'][0]==feature].reset_index()        
        ax.plot(gamdata_plotdataX.loc[0,'x'], gamdata_plotdataX.loc[0,'fit']+gfc[gfc[0]==feature]['intercept_spline'].iloc[0], linewidth=5, color='b', label='spline')
        

        if ext_obj1 is not None:
            gamdata_plotdataX1 = ext_obj1.gamdata_plotdata1d[label][ext_obj1.gamdata_plotdata1d[label][0]==feature].reset_index()                    
            ax.plot(gamdata_plotdataX1.loc[0,'x'], gamdata_plotdataX1.loc[0,'fit'], linewidth=5, color='r', label='interaction')
        
        if ext_obj2 is not None:
            gamdata_plotdataX2 = ext_obj2.gamdata_plotdata1d[label][ext_obj2.gamdata_plotdata1d[label][0]==feature].reset_index()                    
            ax.plot(gamdata_plotdataX2.loc[0,'x'], gamdata_plotdataX2.loc[0,'fit'], linewidth=5, color='r')
        
        for i in raw_fits.index:            
            ax.plot(spline_x, raw_fits.loc[i,0](spline_x), alpha=alpha)            
            # b1 = self.plot_range['Feature'] == raw_fits.loc[i,'Feature']
            # b2 = self.plot_range['site_d'] == raw_fits.loc[i,'site_d']
            # b3 = self.plot_range['site_m'] == raw_fits.loc[i,'site_m']
            # minx = self.plot_range[np.logical_and(b1,np.logical_and(b2,b3))].iloc[0,3]
            # maxx = self.plot_range[np.logical_and(b1,np.logical_and(b2,b3))].iloc[0,4]            
            # xcood = np.array(gamdata_plotdataX.loc[0,'x'])
            # xcood = xcood[xcood>=minx]
            # xcood = xcood[xcood<=maxx]
#            ax.plot(xcood, raw_fits.loc[i,0](xcood), alpha=alpha)
            
        #ax.scatter(shapdf.loc[:,'Name'], shapdf.loc[:,'val'])
        ylimmax = (gamdata_plotdataX.loc[0,'fit']+gfc[gfc[0]==feature]['intercept_spline'].iloc[0]).max()*1.1
        ylimmin = (gamdata_plotdataX.loc[0,'fit']+gfc[gfc[0]==feature]['intercept_spline'].iloc[0]).min()*1.1
        
#        ax.set_ylim([ylimmin, ylimmax])
        
        ax.fill_between(gamdata_plotdataX.loc[0,'x'], gamdata_plotdataX.loc[0,'fit']-1.96*np.array(gamdata_plotdataX.loc[0,'se']), gamdata_plotdataX.loc[0,'fit']+1.96*np.array(gamdata_plotdataX.loc[0,'se']), alpha=0.3)  
        print((min(gamdata_plotdataX.loc[0,'x']), max(gamdata_plotdataX.loc[0,'x'])))
        ax.set_xlim(min(gamdata_plotdataX.loc[0,'x']), max(gamdata_plotdataX.loc[0,'x']))
        ax.grid(alpha=0.2)
        pic_filname = self.filename+'_'+feature+'.svg'
        pic_filname = pic_filname.replace('::','_').replace('/','per').replace('(','_').replace(')','_')
        ax.set_title(feature)
        if axis_label:
            ax.set_xlabel('value')
            ax.set_ylabel('log_odd_change')
        else:
            ax.set_xlabel('')
            ax.set_ylabel('')            
        ax.legend()
        
        
        ymax = max(gamdata_plotdataX.loc[0,'fit']+gfc[gfc[0]==feature]['intercept_spline'].iloc[0])
        ymin = min(gamdata_plotdataX.loc[0,'fit']+gfc[gfc[0]==feature]['intercept_spline'].iloc[0])
        # Calculate the range and margin
        yrange = ymax - ymin

        # Set the new limits with margin
        ymin_with_margin = ymin - margin/2 * yrange
        ymax_with_margin = ymax + margin/2 * yrange
        
        ax.set_ylim(ymin_with_margin, ymax_with_margin)

#        plt.savefig(pic_filname, bbox_inches ='tight')
#        plt.show()
        return ylimmin, ylimmax

    def plot_meta1_cat(self, feature, ax=None, axis_label=True):
        if ax is None:
            ax = plt.gca()
        gfc = self.gen_fit_character()
        
#        dfmta = self.gamdata_plotdatap['linear'][self.gamdata_plotdatap['linear'][0]==feature].reset_index().loc[0,3] 

        # sitecatshap = self.shapdf[self.shapdf['Feature']==feature].dropna()[['site_d', 'site_m', 'val', 'Name']].groupby(['site_d', 'site_m', 'Name']).agg(['count', np.mean, np.std, np.max, np.min]).reset_index()
    
        columns_select = [feature+'_Names', feature+'_vals', 'site_m', 'site_d']
        shapdf = self.shapdf[columns_select].dropna()
        shapdf.columns = ['Name', 'val', 'site_m', 'site_d']
        shapdf['Feature'] = feature
        shapdf = shapdf.sort_values('Name')

        sitecatshap = shapdf.dropna()[['site_d', 'site_m', 'val', 'Name']].groupby(['site_d', 'site_m', 'Name']).agg(['count', np.mean, np.std, np.max, np.min]).reset_index()
    
        sitecatshap.columns = ['site_d', 'site_m', 'Name','count', 'mean', 'std', 'max', 'min']
        sitecatshap['ste'] = sitecatshap['std']/sitecatshap['count']
        sitecatshap['sitemsited'] = sitecatshap['site_d']+':'+sitecatshap['site_m']

        df = sitecatshap
        df['Name'] = df['Name'].astype(str)

        im = sns.boxplot(y=df["mean"], x=df["Name"], ax=ax)
        ax.axhline(y=0, linestyle='--', color='r')
        # ax.axhline(y=dfmta.loc[0,'y']+gfc[gfc[0]==feature]['intercept_linear'].iloc[0], linestyle='--', color='cyan')
        # ax.axhline(y=dfmta.loc[1,'y']+gfc[gfc[0]==feature]['intercept_linear'].iloc[0], linestyle='--', color='gold')
        
        ax.axhline(y=gfc[gfc[0]==feature]['intercept_linear'].iloc[0], linestyle='--', color='cyan')
        ax.axhline(y=gfc[gfc[0]==feature]['slope_linear'].iloc[0]+gfc[gfc[0]==feature]['intercept_linear'].iloc[0], linestyle='--', color='gold')
        
        ax.set_title(feature)
        if axis_label:
            ax.set_xlabel('exists')
            ax.set_ylabel('mean_log_odd_change')
        else:
            ax.set_xlabel('')
            ax.set_ylabel('')    
            
        return im


    def plot_meta1_all(self, filter_sitem=[], n_jobs=4, ext_obj1=None, ext_obj2=None, label='spline'):
        Parallel(n_jobs=n_jobs)(delayed(self.plot_meta1)(feature, filter_sitem=['MCRI'], ext_obj1=ext_obj1, ext_obj2=ext_obj2) for feature in self.shapdf['Feature'].unique())
            
    def copy(self, ext_obj):
        self.fits = ext_obj.fits
        self.gamdata_fitdata = ext_obj.gamdata_fitdata
        self.gamdata_plotdata1d = ext_obj.gamdata_plotdata1d
        self.gamdata_plotdata2d = ext_obj.gamdata_plotdata2d
        self.gamdata_plotdatag = ext_obj.gamdata_plotdatag
        self.fits = ext_obj.fits
        self.shapdf = ext_obj.shapdf
        self.plot_range = ext_obj.plot_range
        self.features = ext_obj.features
        
    def save(self):
        with open(self.home_data_directory+self.filename+'.pkl', 'wb') as f:
            pickle.dump(self, f, protocol=4)        
            
    def fit_table(self):
        self.gamdata_fitdata['AUC_slope'] = self.gamdata_fitdata['p.coeff'].apply(lambda x: x[-1])
        self.gamdata_fitdata['r.sq'] = self.gamdata_fitdata['r.sq'].apply(lambda x: x[-1])
        if -1 in self.gamdata_fitdata.columns:
            self.fit_table_df = self.gamdata_fitdata[[-1, 0, 'AUC_slope', 'r.sq']]
        else:
            self.fit_table_df = self.gamdata_fitdata[[0, 'AUC_slope', 'r.sq']]

            
            
            
    def cal_plot_range(self):
#        self.plot_range = self.shapdf[['Feature', 'site_d', 'site_m', 'Name']].groupby(['Feature', 'site_d', 'site_m']).agg([min,max]).reset_index()
        xxx = self.shapdf[['site_d', 'site_m'] + [x+"_Names" for x in self.features if x not in list(self.cattarget)]]
        xxx.columns = ['site_d', 'site_m'] + [x for x in self.features]
        xxx = xxx.groupby(['site_d', 'site_m']).agg([min,max]).reset_index()
        xxx = xxx.rename({'ORIGINAL_BMI':'ORIGINALBMI'},axis=1)

        object_columns = [col for col in xxx.columns if xxx[col].dtype == 'float']
        columns_to_keep = [('site_d',""), ('site_m',"")] + object_columns

        xxx = xxx[columns_to_keep]

        # Step 1: Reset the MultiIndex
        df_reset = xxx.copy()
        # Step 2: Rename the columns
        df_reset.columns = ['site_d', 'site_m'] + [f'{col[0]}_{col[1]}' if col[1] else col[0] for col in df_reset.columns[2:]]
        # Step 3: Melt the DataFrame
        df_melted = df_reset.melt(id_vars=['site_d', 'site_m'], var_name='Feature', value_name='Value')
        # Step 4: Split the 'Feature' column to separate 'Feature' and 'Min/Max'
        df_melted[['Feature', 'Min/Max']] = df_melted['Feature'].str.split('_', expand=True)
        # Step 5: Pivot the DataFrame to have separate columns for 'min' and 'max'
        df_pivot = df_melted.pivot_table(index=['site_d', 'site_m', 'Feature'], columns='Min/Max', values='Value').reset_index()
        # Step 6: Remove '_Names' from 'Feature' column
        df_pivot['Feature'] = df_pivot['Feature'].str.replace('_Names', '')

        self.plot_range = df_pivot

    def plot_single_site_feature(self, site_d, site_m, feature, label='spline'):
        gamdata_plotdataX = self.gamdata_plotdata1d[label][self.gamdata_plotdata1d[label][0]==feature].reset_index()
        shapdf = self.shapdf[self.shapdf['Feature'] == feature]
        shapdf = shapdf[shapdf['site_d'] == site_d]
        shapdf = shapdf[shapdf['site_m'] == site_m]
        shapdf = shapdf.sort_values('Name')
        raw_fits = self.fits[self.fits[0]==feature].iloc[0,1]
        raw_fits = raw_fits[raw_fits['site_d'] == site_d]
        raw_fits = raw_fits[raw_fits['site_m'] == site_m]

        formula = interpolate.UnivariateSpline(shapdf['Name'], shapdf['val'],k=3)
         
        plt.scatter(shapdf['Name'], shapdf['val'])
        
        b1 = self.plot_range['Feature'] == raw_fits.iloc[0,:]['Feature']
        b2 = self.plot_range['site_d'] == raw_fits.iloc[0,:]['site_d']
        b3 = self.plot_range['site_m'] == raw_fits.iloc[0,:]['site_m']
        minx = self.plot_range[np.logical_and(b1,np.logical_and(b2,b3))].iloc[0,3]
        maxx = self.plot_range[np.logical_and(b1,np.logical_and(b2,b3))].iloc[0,4]            
        xcood = np.array(gamdata_plotdataX.loc[0,'x'])
        xcood = xcood[xcood>=minx]
        xcood = xcood[xcood<=maxx]        
        
        plt.plot(gamdata_plotdataX['x'][0], gamdata_plotdataX['fit'][0], 'g')
        plt.plot(xcood, raw_fits.iloc[0,:][0](xcood), 'y')
        plt.plot(xcood, formula(xcood), 'r')

    def plot_meta2d_sub(self, x, y, z, xlab, ylab, cmap='bwr', interpolation='lanczos', ax = None):
        #plt.imshow(np.flipud(np.array(plotshap.gamdata_plotdata2d.iloc[0,:]['fit']).reshape(40,40)))
        if ax is None:
            ax = plt.gca()         
        endpt = max(max(z), abs(min(z)))
        extent=[min(x),max(x),min(y),max(y)]
        z = np.flipud(z.reshape(len(x),len(x)))
        im = ax.imshow(z, cmap=cmap, interpolation=interpolation, vmax=endpt, vmin=-endpt, extent=extent, aspect='auto')
#        ax.colorbar()
        ax.set_xlabel(xlab)
        ax.set_ylabel(ylab)
#        plt.show()
        return im

    def plot_meta2_cat(self, f1, f2, ext_obj=None, ax=None, verbose=True, legend_size=8, mode2='nouni'):       
        if f1 in list(self.cattarget):
            ft = f2
            f2 = f1
            f1 = ft
        gamdata_plotdataX = self.gamdata_plotdata1d['double_interaction'].copy()
        gamdata_plotdataX = gamdata_plotdataX[((gamdata_plotdataX[-1]==f1) & (gamdata_plotdataX[0]==f2)) | ((gamdata_plotdataX[-1]==f2) & (gamdata_plotdataX[0]==f1))]        
        # gamdata_plotdatayb = gamdata_plotdataX[~(gamdata_plotdataX['ylab'].str.contains('y0|y1'))].reset_index(drop=True)
        # gamdata_plotdatay0 = gamdata_plotdataX[gamdata_plotdataX['ylab'].str.contains('y0')].reset_index(drop=True)
        # gamdata_plotdatay1 = gamdata_plotdataX[gamdata_plotdataX['ylab'].str.contains('y1')].reset_index(drop=True)        
        
        gamdata_plotdatayb = gamdata_plotdataX[~(gamdata_plotdataX['ylab'].str.contains('FALSE|TRUE'))].reset_index(drop=True)
        gamdata_plotdatay0 = gamdata_plotdataX[gamdata_plotdataX['ylab'].str.contains('FALSE')].reset_index(drop=True)
        gamdata_plotdatay1 = gamdata_plotdataX[gamdata_plotdataX['ylab'].str.contains('TRUE')].reset_index(drop=True)        
        
        all_intercept = self.gamdata_fitdata['double_interaction'][((self.gamdata_fitdata['double_interaction'][-1]==f1) & (self.gamdata_fitdata['double_interaction'][0]==f2)) | ((self.gamdata_fitdata['double_interaction'][-1]==f2) & (self.gamdata_fitdata['double_interaction'][0]==f1))]['p.coeff'].iloc[0][0]
        y_intercept = self.gamdata_fitdata['double_interaction'][((self.gamdata_fitdata['double_interaction'][-1]==f1) & (self.gamdata_fitdata['double_interaction'][0]==f2)) | ((self.gamdata_fitdata['double_interaction'][-1]==f2) & (self.gamdata_fitdata['double_interaction'][0]==f1))]['p.coeff'].iloc[0][1]

        if ax is None:
            ax = plt.gca() 
        ax.axhline(y=0, linestyle='--', color='r')
        if mode2 == 'all'or mode2=='nouni':
            im = ax.plot(gamdata_plotdatay0.loc[0,'x'], all_intercept+gamdata_plotdatay0.loc[0,'fit']+gamdata_plotdatayb.loc[0,'fit'], linewidth=5, color='y', label='Binary=False')
            # ax.fill_between(gamdata_plotdatay0.loc[0,'x'], gamdata_plotdatay0.loc[0,'fit']+gamdata_plotdatayb.loc[0,'fit']-1.96*np.array(gamdata_plotdatay0.loc[0,'se']), 
            #                                                 gamdata_plotdatay0.loc[0,'fit']+gamdata_plotdatayb.loc[0,'fit']+1.96*np.array(gamdata_plotdatay0.loc[0,'se']), alpha=0.3)  

            ax.plot(gamdata_plotdatay1.loc[0,'x'], all_intercept+y_intercept+gamdata_plotdatay1.loc[0,'fit']+gamdata_plotdatayb.loc[0,'fit'], linewidth=5, color='b', label='Binary=True')
            # ax.fill_between(gamdata_plotdatay1.loc[0,'x'], gamdata_plotdatay1.loc[0,'fit']+gamdata_plotdatayb.loc[0,'fit']-1.96*np.array(gamdata_plotdatay1.loc[0,'se']), 
            #                                                 gamdata_plotdatay1.loc[0,'fit']+gamdata_plotdatayb.loc[0,'fit']+1.96*np.array(gamdata_plotdatay1.loc[0,'se']), alpha=0.3)  
            
        if mode2 == 'interact':
            gfc = ext_obj.gen_fit_character()
            gamdata_plotdataXs = ext_obj.gamdata_plotdata1d['spline'][ext_obj.gamdata_plotdata1d['spline'][0]==f1].reset_index()
            fit0 = interp1d(gamdata_plotdatay0.loc[0,'x'], gamdata_plotdatay0.loc[0,'fit'], kind='quadratic', fill_value='linear')
            fit1 = interp1d(gamdata_plotdatay1.loc[0,'x'], gamdata_plotdatay1.loc[0,'fit'], kind='quadratic', fill_value='linear')
            fit11 = fit1(gamdata_plotdataXs.loc[0,'x'])
            fit00 = fit0(gamdata_plotdataXs.loc[0,'x'])
            # im = ax.plot(gamdata_plotdataXs.loc[0,'x'], gamdata_plotdataXs.loc[0,'fit']+gfc[gfc[0]==f1]['intercept_spline'].iloc[0]+fit11, linewidth=5, marker='o', label='Binary=True', color='b')
            # ax.plot(gamdata_plotdataXs.loc[0,'x'], gamdata_plotdataXs.loc[0,'fit']+gfc[gfc[0]==f1]['intercept_spline'].iloc[0]+fit00, linewidth=5, marker='o', label='Binary=False', color='y')         
            im = ax.plot(gamdata_plotdataXs.loc[0,'x'], y_intercept+fit11-fit00, linewidth=5, marker='o', label='Interaction', color='b')
            
        if ext_obj is not None and mode2== 'all':
            gfc = ext_obj.gen_fit_character()
            gamdata_plotdataXs = ext_obj.gamdata_plotdata1d['spline'][ext_obj.gamdata_plotdata1d['spline'][0]==f1].reset_index()        
            ax.plot(gamdata_plotdataXs.loc[0,'x'], gamdata_plotdataXs.loc[0,'fit']+gfc[gfc[0]==f1]['intercept_spline'].iloc[0], linewidth=5, marker='o', label='Univariate')
                # ax.fill_between(gamdata_plotdataXs.loc[0,'x'], gamdata_plotdataXs.loc[0,'fit']+gfc[gfc[0]==f1]['intercept_spline'].iloc[0]-1.96*np.array(gamdata_plotdataXs.loc[0,'se']), gamdata_plotdataXs.loc[0,'fit']+gfc[gfc[0]==f1]['intercept_spline'].iloc[0]+1.96*np.array(gamdata_plotdataXs.loc[0,'se']), alpha=0.3) 

      
            
        ax.set_xlim(min(gamdata_plotdatay0.loc[0,'x']+gamdata_plotdatay1.loc[0,'x']), max(gamdata_plotdatay0.loc[0,'x']+gamdata_plotdatay1.loc[0,'x']))
        ax.grid(alpha=0.2)
        #pic_filname = self.filename+'_'+feature+'.svg'
        #pic_filname = pic_filname.replace('::','_').replace('/','per').replace('(','_').replace(')','_')
        if verbose:
            ax.set_title(gamdata_plotdatay0.loc[0,-1]+':'+gamdata_plotdatay0.loc[0,0])
            ax.set_xlabel('value')
            ax.set_ylabel('log_odd_change')
        ax.legend(fontsize=legend_size)     
        return im

    def plot_meta2_cont(self, f1, f2, ext_obj=None, ax=None, verbose=True, alpha=0.3, plottype='full', contour=False, fontsize=8, ticksize=8):

#        print(f1,f2)
        rea2d = self.gamdata_plotdata2d['double_interaction'][self.gamdata_plotdata2d['double_interaction']['ylab']=='Name.y']
        def converfit(df):
            return [float(x) if x !='NA' else np.nan for x in df]
        rea2d['numfit'] = rea2d['fit'].map(converfit)
        rea2d['fitmax'] = rea2d['numfit'].apply(np.nanmax)
        rea2d['fitmin'] = rea2d['numfit'].apply(np.nanmin)
        rea2d['fitmaxmin'] = rea2d['fitmax']-rea2d['fitmin']
        rea2d= rea2d.sort_values('fitmaxmin', ascending=False)
        def applymean(df):
            return np.nanmean(np.abs(np.array(df)))
        rea2d['fitmean'] = rea2d['numfit'].apply(applymean)
        rea2d= rea2d.sort_values('fitmean', ascending=False)
        rea2d = pd.concat([rea2d[np.logical_and(rea2d[-1]==f1, rea2d[0]==f2)], rea2d[np.logical_and(rea2d[-1]==f2, rea2d[0]==f1)]])

        if rea2d[-1].iloc[0]==f1:
            xlab = rea2d.iloc[0][-1]
            ylab = rea2d.iloc[0][0]
            x = rea2d.iloc[0]['x']
            y = rea2d.iloc[0]['y']
            z = np.array(rea2d.iloc[0]['numfit'])
            zxycoor = z.reshape(len(x),len(x))
            X, Y = np.meshgrid(x, y)
        else:
            xlab = rea2d.iloc[0][0]
            ylab = rea2d.iloc[0][-1]
            x = rea2d.iloc[0]['y']
            y = rea2d.iloc[0]['x']
            z = np.array(rea2d.iloc[0]['numfit'])
            zxycoor = z.reshape(len(x),len(x)).T
            X, Y = np.meshgrid(x, y)
        
        gamdata_plotdataX = self.gamdata_plotdata1d['double_interaction'].copy()
        gamdata_plotdataX = gamdata_plotdataX[((gamdata_plotdataX[-1]==f1) & (gamdata_plotdataX[0]==f2)) | ((gamdata_plotdataX[-1]==f2) & (gamdata_plotdataX[0]==f1))]        
        
        if gamdata_plotdataX.shape[0] == 4:
            gamdata_plotdataX = gamdata_plotdataX[((gamdata_plotdataX[-1]==f1) & (gamdata_plotdataX[0]==f2))]    
        
        all_intercept = self.gamdata_fitdata['double_interaction'][((self.gamdata_fitdata['double_interaction'][-1]==f1) & (self.gamdata_fitdata['double_interaction'][0]==f2)) | ((self.gamdata_fitdata['double_interaction'][-1]==f2) & (self.gamdata_fitdata['double_interaction'][0]==f1))]['p.coeff'].iloc[0][0]

        if rea2d[-1].iloc[0]==f1:
            gamdata_plotdatax0 = gamdata_plotdataX[gamdata_plotdataX['xlab'].str.contains('Name.x')].reset_index(drop=True)
            gamdata_plotdatay1 = gamdata_plotdataX[gamdata_plotdataX['xlab'].str.contains('Name.y')].reset_index(drop=True)        
        else:
            gamdata_plotdatay1 = gamdata_plotdataX[gamdata_plotdataX['xlab'].str.contains('Name.x')].reset_index(drop=True)
            gamdata_plotdatax0 = gamdata_plotdataX[gamdata_plotdataX['xlab'].str.contains('Name.y')].reset_index(drop=True)        
            
        gamdata_plotdataX = gamdata_plotdataX[((gamdata_plotdataX[-1]==f1) & (gamdata_plotdataX[0]==f2)) | ((gamdata_plotdataX[-1]==f2) & (gamdata_plotdataX[0]==f1))]        
        
        
            
        
        
        f = interp1d(gamdata_plotdatax0.loc[0,'x'], gamdata_plotdatax0.loc[0,'fit'], kind='quadratic', fill_value='linear')
        newz_x0 = f(x)
        f = interp1d(gamdata_plotdatay1.loc[0,'x'], gamdata_plotdatay1.loc[0,'fit'], kind='quadratic', fill_value='linear')
        newz_y1 = f(y)          

        
        newz_X0, newz_Y1 = np.meshgrid(newz_x0, newz_y1)
#        zxycoor = z.reshape(len(x),len(x)).T
#        zxycoor_sum=zxycoor+all_intercept+newz_y1[:,np.newaxis]+newz_x0[np.newaxis,:]
        zxycoor_sum=zxycoor+all_intercept+newz_Y1+newz_X0


    
    
        #Subtract f2 contribution
        gfc = ext_obj.gen_fit_character()
        gamdata_plotdataXs = ext_obj.gamdata_plotdata1d['spline'][ext_obj.gamdata_plotdata1d['spline'][0]==f2].reset_index()        
        f = interp1d(gamdata_plotdataXs.loc[0,'x'], gamdata_plotdataXs.loc[0,'fit']+gfc[gfc[0]==f2]['intercept_spline'].iloc[0], kind='quadratic', fill_value='extrapolate')
        newz_y2 = f(y)
        
        newz_X0, newz_Y2 = np.meshgrid(newz_x0, newz_y2)
        zxycoor_sum2 = zxycoor_sum-newz_Y2       
#        zxycoor_sum2 = zxycoor_sum-newz_y2[:,np.newaxis]
        
        #Interaction only
        zxycoor_sumo=zxycoor
        
        #Interaction + f1 univariate
        gfc = ext_obj.gen_fit_character()
        gamdata_plotdataXs = ext_obj.gamdata_plotdata1d['spline'][ext_obj.gamdata_plotdata1d['spline'][0]==f1].reset_index()        
        f = interp1d(gamdata_plotdataXs.loc[0,'x'], gamdata_plotdataXs.loc[0,'fit']+gfc[gfc[0]==f1]['intercept_spline'].iloc[0], kind='quadratic', fill_value='extrapolate')
        newz_x3 = f(x)       
        
        newz_X3, newz_Y2 = np.meshgrid(newz_x3, newz_y2)
        zxycoor_sumf1 = zxycoor+newz_X3   
#        zxycoor_sumf1=zxycoor+newz_y3[np.newaxis,:]
        
        #Interaction only normalized
        gfc = ext_obj.gen_fit_character()
        gamdata_plotdataXs = ext_obj.gamdata_plotdata1d['spline'][ext_obj.gamdata_plotdata1d['spline'][0]==f1].reset_index()        
        f = interp1d(gamdata_plotdataXs.loc[0,'x'], gamdata_plotdataXs.loc[0,'fit']+gfc[gfc[0]==f1]['intercept_spline'].iloc[0], kind='quadratic', fill_value='extrapolate')       
        newz_X3 = f(x)       

        newz_X3, newz_Y2 = np.meshgrid(newz_x3, newz_y2)
        zxycoor_sum2 = zxycoor/newz_X3           
#        zxycoor_sumon=zxycoor/newz_y3[np.newaxis,:]
        
        if ax is None:
            ax = plt.gca()   
            
        if not contour:
            ax.axhline(y=0, linestyle='--', color='r')
        
        # Create a color map
        cmap = plt.cm.get_cmap('jet')  # You can choose any available colormap

        # Normalize the y values to use with the colormap
        if not contour:
            norm = plt.Normalize(min(y), max(y))


        # Loop through each y value
        for i in range(len(y)):
            # Extract the z values corresponding to this y
            if plottype=='full':
                z_values = zxycoor_sum[i,:]
            elif plottype=='subtract':
                z_values = zxycoor_sum2[i,:]
            elif plottype=='interactonly':
                z_values = zxycoor_sumo[i,:]
            elif plottype=='addition':
                z_values = zxycoor_sumf1[i,:]
            
            # Plot x vs z, coloring by y value
            if not contour:
                ax.plot(x, z_values, color=cmap(norm(y[i])), label=f'y = {y[i]}', alpha=alpha)
            else:
                contour = ax.contour(X, Y, zxycoor, cmap='viridis', levels=10)
                ax.clabel(contour, inline=True, fontsize=ticksize)
                
        if ext_obj is not None and plottype!='interactonly':
            gfc = ext_obj.gen_fit_character()
            gamdata_plotdataXs = ext_obj.gamdata_plotdata1d['spline'][ext_obj.gamdata_plotdata1d['spline'][0]==f1].reset_index()        
            if not contour:
                ax.plot(gamdata_plotdataXs.loc[0,'x'], gamdata_plotdataXs.loc[0,'fit']+gfc[gfc[0]==f1]['intercept_spline'].iloc[0], linewidth=5, label='Univariate', color='g')    

            
        # Create a color bar
        if not contour:
            sm = plt.cm.ScalarMappable(cmap=cmap, norm=norm)
            sm.set_array([])
            # Creating a color bar in the specified subplot
            cb = plt.colorbar(sm, ax=ax)
            cb.set_label(self.translator.custom_translate_omop_2022_2_fig2(f2), fontsize=fontsize) 
            cb.ax.tick_params(labelsize=ticksize)
        
        if verbose:
            ax.set_xlabel(f1)
            if not contour:
                ax.set_ylabel('SHAP', labelpad=fontsize)
            ax.set_title(f2)
            
    def plot_meta2(self, f1, f2, ax=None):
        rea2d = self.gamdata_plotdata2d['double_interaction'][self.gamdata_plotdata2d['double_interaction']['ylab']=='Name.y']
        def converfit(df):
            return [float(x) if x !='NA' else np.nan for x in df]
        rea2d['numfit'] = rea2d['fit'].map(converfit)
        rea2d['fitmax'] = rea2d['numfit'].apply(np.nanmax)
        rea2d['fitmin'] = rea2d['numfit'].apply(np.nanmin)
        rea2d['fitmaxmin'] = rea2d['fitmax']-rea2d['fitmin']
        rea2d= rea2d.sort_values('fitmaxmin', ascending=False)
        def applymean(df):
            return np.nanmean(np.abs(np.array(df)))
        rea2d['fitmean'] = rea2d['numfit'].apply(applymean)
        rea2d= rea2d.sort_values('fitmean', ascending=False)
        rea2d = pd.concat([rea2d[np.logical_and(rea2d[-1]==f1, rea2d[0]==f2)], rea2d[np.logical_and(rea2d[-1]==f2, rea2d[0]==f1)]])
        xlab = rea2d.iloc[0][-1]
        ylab = rea2d.iloc[0][0]
        x = rea2d.iloc[0]['x']
        y = rea2d.iloc[0]['y']
        z = np.array(rea2d.iloc[0]['numfit'])
        return self.plot_meta2d_sub(x, y, z, xlab, ylab, ax=ax)
                           
    def combine_interaction_json(self, path='/home/hoyinchan/blue/program_data/AKI_CDM_PY/MetaRegression/gam2d_tmp/'):
        alliles = [x for x in os.listdir(path) if 'gam2d_tmp' in x]
        feature1 = list()
        feature2 = list()
        fitdata = list()
        plotdata = list()
        for file in alliles:
            with open(path+file, 'r') as infile:
                tmp = json.load(infile)
                feature1.append(tmp[0])
                feature2.append(tmp[1])
                fitdata.append(tmp[2])
                plotdata.append(tmp[3])
        pd.DataFrame({'feature1':feature1, 'feature2':feature2, 'fitdata':fitdata, 'plotdata':plotdata} ).to_pickle('/home/hoyinchan/blue/program_data/AKI_CDM_PY/MetaRegression/gamalltmp_double_interaction_noAUC_json.pkl')
        
    def generate_single_table(self):
        tmp_tt = self.fit_table_df.merge(plotshapsi.fit_table_df, on=[0], how='inner')
        tmp_tt.columns = ['Feature', 'AUC_slope_no_int', 'r^2_no_int', 'AUC_slope_int', 'r^2_int']
        tmp_tt = tmp_tt[['Feature', 'AUC_slope_no_int', 'AUC_slope_int', 'r^2_no_int', 'r^2_int']]
        tmp_tt['LOINC_NUM']=tmp_tt['Feature']
        tmp_tt['LOINC_NUM'] = tmp_tt['LOINC_NUM'].str.split(':').str[2].str.split('(').str[0]
        tmp_tt = tmp_tt.merge(loinctranslate[['LOINC_NUM', 'LONG_COMMON_NAME']], on='LOINC_NUM', how='left').drop('LOINC_NUM',axis=1)
        return np.round(tmp_tt,2)
    
    def generate_double_table(self):
        xx2 = self.fit_table_df.sort_values('r.sq', ascending=False)
        xx2 = xx2[np.logical_and(xx2[-1]!='LAB::33037-3(mmol/L)', xx2[0]!='LAB::33037-3(mmol/L)')]
        xx2 = xx2.merge(plotshapsi.fit_table_df, left_on = [-1], right_on=[0], how='left').merge(plotshapsi.fit_table_df, left_on = ['0_x'], right_on=[0], how='left')
        xx2 = xx2.drop([0, 'AUC_slope_x', '0_y', 'AUC_slope_y', 'AUC_slope'],axis=1)
        xx2.columns = ['Feature1', 'Feature2', 'r.sq', 'r.sq_f1', 'r.sq_f2']
        xx2['r.sq_max'] = xx2[[ 'r.sq_f1', 'r.sq_f2']].max(axis=1)
        xx2['r.sq_diff'] = xx2['r.sq']-xx2['r.sq_max']
        xx2 = xx2.sort_values('r.sq_diff',ascending=False).head(18)
        xx2['LOINC_NUM']=xx2['Feature1']
        xx2['LOINC_NUM'] = xx2['LOINC_NUM'].str.split(':').str[2].str.split('(').str[0]
        xx2 = xx2.merge(loinctranslate[['LOINC_NUM', 'LONG_COMMON_NAME']], on='LOINC_NUM', how='left').drop('LOINC_NUM',axis=1)
        xx2['LOINC_NUM']=xx2['Feature2']
        xx2['LOINC_NUM'] = xx2['LOINC_NUM'].str.split(':').str[2].str.split('(').str[0]
        xx2 = xx2.merge(loinctranslate[['LOINC_NUM', 'LONG_COMMON_NAME']], on='LOINC_NUM', how='left').drop('LOINC_NUM',axis=1)
        xx2 = xx2.drop(['r.sq_max','r.sq_diff'],axis=1)
        xx2.columns = ['Feature1', 'Feature2', 'r.sq', 'r.sq_f1', 'r.sq_f2', 'Name1', 'Name2']
        return xx2
    
    def get_raw_stat(self):
        sites = ['MCW', 'UIOWA', 'UMHC', 'UNMC', 'UofU', 'UTHSCSA', 'UTSW', 'IUR', 'KUMC', 'UPITT']
        data_onset = list()
        data_demo = list()
        for site in sites:
            onset = pd.read_pickle('/home/hoyinchan/blue/Data/data2021/data2021/'+site+'/p0_onset_'+site+'.pkl')
            if site == 'IUR':
                onset = onset[onset['ADMIT_DATE'].str.split('-').str[0].astype(int)<=2016]
            demo = pd.read_pickle('/home/hoyinchan/blue/Data/data2021/data2021/'+site+'/p0_demo_'+site+'.pkl')
            demo = demo.merge(onset[['PATID', 'ENCOUNTERID']], how='inner')
            onset['site'] = site
            demo['site'] = site
            data_onset.append(onset)
            data_demo.append(demo)
        data_onsets = pd.concat(data_onset)
        data_demos = pd.concat(data_demo)

        def nullcounter(data_onsets, feature):
            NONAKI_COUNT = data_onsets[['site', feature]].copy()
            NONAKI_COUNT[feature] = NONAKI_COUNT[feature].notnull()
            NONAKI_COUNT = NONAKI_COUNT.groupby('site').sum().reset_index()
            return NONAKI_COUNT

        PATID_COUNT = data_onsets[['site', 'PATID']].drop_duplicates().groupby('site').count().reset_index()
        ENCOUNTERID_COUNT = data_onsets[['site', 'ENCOUNTERID']].drop_duplicates().groupby('site').count().reset_index()
        NONAKI_COUNT = nullcounter(data_onsets, 'NONAKI_ANCHOR')
        AKI1_COUNT = nullcounter(data_onsets, 'AKI1_ONSET')
        AKI2_COUNT = nullcounter(data_onsets, 'AKI2_ONSET')
        AKI3_COUNT = nullcounter(data_onsets, 'AKI3_ONSET')

        onet_stat = PATID_COUNT.merge(ENCOUNTERID_COUNT, on='site').merge(NONAKI_COUNT, on='site').merge(AKI1_COUNT, on='site').merge(AKI2_COUNT, on='site').merge(AKI3_COUNT, on='site')

        onet_stat.columns = ['site', 'Patients', 'Encounter', 'Non-AKI', 'AKI1', 'AKI2', 'AKI3']

        AGE_COUNT = np.round(data_demos[['site', 'AGE']].drop_duplicates().groupby('site').agg([np.mean, np.std]).reset_index(),2)
        AGE_COUNT.columns = ['site', 'x', 'y']
        AGE_COUNT['Age'] = AGE_COUNT['x'].astype(str) + '(' + AGE_COUNT['y'].astype(str) + ')'
        AGE_COUNT=AGE_COUNT.drop(['x', 'y'], axis=1)

        SEX_COUNT = data_demos[['site', 'SEX', 'PATID']].groupby(['site', 'SEX']).count().reset_index()
        SEX_COUNT = SEX_COUNT[(SEX_COUNT['SEX']=='M') | (SEX_COUNT['SEX']=='F')]
        SEX_COUNT = SEX_COUNT.pivot(index='site', columns='SEX', values='PATID').astype(int).reset_index()
        SEX_COUNT.columns = ['site', 'Female', 'Male']
        RACE_COUNT = data_demos[['site', 'RACE', 'PATID']].groupby(['site', 'RACE']).count().reset_index()
        RACE_COUNT = RACE_COUNT[[True if x in ['01', '02', '03', '04', '05'] else False for x in RACE_COUNT['RACE']]]
        RACE_COUNT = RACE_COUNT.pivot(index='site', columns='RACE', values='PATID').fillna(0).astype(int).reset_index()
        RACE_COUNT.columns = ['site', 'American Indian or Alaska Native', 'Asian', 'Black or African American', 'Native Hawaiian or Other Pacific Islander', 'White']

        demo_stat = AGE_COUNT.merge(SEX_COUNT, on='site').merge(RACE_COUNT, on='site')

        return onet_stat.merge(demo_stat, on='site').T

    def plot_fig2(self, select_row, hspace=0.2, wspace=0.2, legend_fontsize=10, nrows=6, ncols=4, margin = 1, namesuffix='', fontsize=12, rescale_y=True):
          
        fig, ax = plt.subplots(nrows, ncols, figsize=(15, 15))

        xxf = select_row
#        xxf = self.features[:nrows*ncols-1]
        # Sorting the array to put PX labels at the back
        px_labels = [label for label in xxf if label.startswith('PX:')]
        non_px_labels = [label for label in xxf if not label.startswith('PX:')]


        # Combining the sorted lists
#        xxf = np.array(non_px_labels + px_labels)

        ## TEST
#        xxf = np.array(non_px_labels)
        self.cattarget =  px_labels

        global_legend = None  # Placeholder for the global legend

        j=0

        for i in range(len(xxf)):       
            print(xxf[i])
    #        Columnwise               
    #        i1 = int(np.floor((i+j)/nrows))
    #        i2 = int((i+j)-np.floor((i+j)/nrows)*nrows)
    #        Rowwise
            i1 = int((i+j)-np.floor((i+j)/ncols)*ncols)
            i2 = int(np.floor((i+j)/ncols))

    #        print(i1, i2)

            if i1 == ncols-1 and i2 == 0:
                j = j+1
                i1 = int((i+j)-np.floor((i+j)/ncols)*ncols)
                i2 = int(np.floor((i+j)/ncols))

    #        print(i1, i2)

            if xxf[i] in list(self.cattarget):
                self.plot_meta1_cat(xxf[i], ax=ax[i2][i1], axis_label=False)
            else:
                ylimmin, ylimmax = self.plot_meta1(xxf[i], filter_sitem=['MCRI'], ax=ax[i2][i1], alpha=0.1, axis_label=False, margin=margin)
                if rescale_y:
                    ax[i2][i1].set_ylim([ylimmin*1.4, ylimmax*1.4])
                xlimmin, xlimmax = self.get_minmax(xxf[i])
                print(xlimmin, xlimmax)
                ax[i2][i1].set_xlim([xlimmin, xlimmax])
                
            # Extract legend from subplot (0,0) and remove legends from each subplot        
            if i1 == 0 and i2 == 0:
                global_legend = ax[i2][i1].get_legend_handles_labels()
    #        ax[i2][i1].legend().set_visible(False)        

            # Extracting original legend location
            original_legend_loc = ax[i2][i1].legend_._loc if ax[i2][i1].legend_ else 'upper left'

            # Adding new legend at the original legend's position
            ax[i2][i1].legend([self.translator.custom_translate_omop_2022_2_fig2(xxf[i])], loc=original_legend_loc, handletextpad=0, handlelength=0)
#            ax[i2][i1].set_title('')  # Remove the subplot title        
            ax[i2][i1].set_title('')  # Remove the subplot title        


        # Adding common X and Y labels      
        fig.text(0.5, 0.1, 'Value', ha='center', va='center', fontsize=fontsize)
        fig.text(0.06, 0.5, 'SHAP', ha='center', va='center', rotation='vertical', fontsize=fontsize)

        # Adding the global legend to the figure
        if global_legend:
            ax[0][ncols-1].legend(*global_legend, loc='center', fontsize=legend_fontsize)
            ax[0][ncols-1].axis('off')

        fig.savefig(f"meta_single{namesuffix}.svg", bbox_inches='tight')        
        
    def plot_fig3(self, ext_obj, suffix=''):
        cin2 = self.get_interaction_stat(ext_obj)
        cin2 = cin2[[-1,'0_x','r.sq_diff']].sort_values('r.sq_diff',ascending=False).head(24)
        cin2 = cin2.reset_index(drop=True)
        fig, ax = plt.subplots(6, 4, figsize=(15, 20))        
        for i in cin2.index:
            i1 = int(np.floor(i/6))
            i2 = int(i-np.floor(i/6)*6)      
            if (cin2.loc[i,-1] in list(self.cattarget)) | (cin2.loc[i,'0_x'] in list(self.cattarget)):
                self.plot_meta2_cat(cin2.loc[i,-1], cin2.loc[i,'0_x'], ax=ax[i2][i1])
            else:
                im = self.plot_meta2(cin2.loc[i,-1], cin2.loc[i,'0_x'], ax=ax[i2][i1])
                plt.colorbar(im, ax=ax[i2,i1])
        fig.savefig(f"meta_double{suffix}.svg", bbox_inches ='tight')

    def plot_fig3_cat(self, ext_obj, numrow, figsize=None, suffix='', font_ratio=1, external_select=None, mode='max_y', mode2='all'):
        # Define sizes for title, xd-label, and ticks
        title_size = 20*font_ratio  # Adjust size as needed
        xlabel_size = 20*font_ratio  # Adjust size as needed
        tick_size = 20*font_ratio  # Adjust size as needed
        legend_size = 20*font_ratio  # Adjust size as needed
        
        if figsize is None:
            figsize = (20, 5*numrow//2)
            
        cin2 = self.get_interaction_stat(ext_obj, mode=mode)            

        if external_select is not None:
            top10f = pd.DataFrame(external_select)
            top10f.columns = ['top10f']
            cin2 = cin2.merge(top10f, left_on=[-1], right_on=['top10f'], how='inner').drop('top10f',axis=1)
        
        # cin2o = cin2.copy()
        # cin2 = cin2.drop(-1,axis=1)
        # # Rename columns: _x to _tmp, _y to _x, and _tmp to _y
        # cin2.columns = [col.replace('_x', '_tmp') for col in cin2.columns]
        # cin2.columns = [col.replace('_y', '_x') for col in cin2.columns]
        # cin2.columns = [col.replace('_tmp', '_y') for col in cin2.columns]
        # # Rename the -1 column to '0_x'
        # cin2[-1] = cin2['0_y'].copy()
        # cin2 = cin2[cin2o.columns]get_cat_interaction_stat
        # cin2 = pd.concat([cin2, cin2o])
        # cin2 = cin2[cin2['r.sq_diff']>0]
        # cin2 = cin2.sort_values('r.sq_diff',ascending=False).groupby(-1).first().reset_index()            
            
        cin2 = cin2[[x in list(self.cattarget) for x in cin2['0_x']]]
        
        xxx = self.get_cat_interaction_feature()
        xxx['0_x'] = xxx.index
        cin2 = cin2[[-1, '0_x', 'r.sq_spline_y']].merge(xxx, on=[-1, '0_x'])
        cin2['r.sq_diff'] = cin2['r.sq']-cin2['r.sq_spline_y']        
        
        cin2 = cin2[cin2['r.sq_diff']>0]
        cin2 = cin2.sort_values('r.sq_diff', ascending=False)
        cin2 = cin2.groupby('0_x').head(numrow).reset_index()
        cin2 = cin2.sort_values(['0_x', 'r.sq_diff'], ascending=False).reset_index(drop=True)

        fig, ax = plt.subplots(numrow//2, 4, figsize=figsize)        
        midpoint = 0.505
        # Add a line at the midpoint. This requires converting the midpoint to the figure's coordinate system.
        line = Line2D([midpoint, midpoint], [0.1, 0.89], transform=fig.transFigure, color="grey", linestyle="--")
        fig.add_artist(line)   
        
        for i in cin2.index:
            i1 = int(np.floor(i/(numrow//2)))
            i2 = int(i-np.floor(i/(numrow//2))*(numrow//2))   
            self.plot_meta2_cat(cin2.loc[i,-1], cin2.loc[i,'0_x'], ext_obj=ext_obj, ax=ax[i2][i1], verbose=False, legend_size=legend_size, mode2=mode2)      

           # Set x_label for each subplot with increased size
            ax[i2][i1].set_xlabel(self.translator.custom_translate_omop_2022_2_fig2(cin2.loc[i, -1]), fontsize=xlabel_size)

            # Set title for the first row subplots with increased size
            if i2 == 0:
                ax[i2][i1].set_title(self.translator.custom_translate_omop_2022_2_fig2(cin2.loc[i, '0_x']), fontsize=title_size)

            # Increase tick size
            ax[i2][i1].tick_params(axis='both', which='major', labelsize=tick_size)
            # Remove legends from all subplots
            ax[i2][i1].legend().set_visible(False)

        # Only enable legend for the top right subplot
        ax[0][-1].legend().set_visible(True)            
        ax[0][-1].legend().get_title().set_fontsize(legend_size)            
        for text in ax[0][-1].legend().get_texts():
            text.set_fontsize(20)  # Increase legend fontsize
        
        # Set a common y_label for all subplots
        fig.text(0.08, 0.5, 'SHAP', va='center', rotation='vertical', fontsize=xlabel_size)

        fig.savefig(f"meta_double_cat{suffix}.svg", bbox_inches ='tight')     
        return fig,ax
            

    def plot_fig3_cont(self, ext_obj, numrow=24, figsize=None, outputname='Double', plottype='full', suffix='', min_r2=0, best_plot=False, external_select=None, ncol=4, contour=False, font_ratio=1, mode='max_y'):

        cin2 = self.get_interaction_stat(ext_obj, mode=mode)
        cin2 = cin2[~cin2['0_x'].str.contains('PX')]
        cin2 = cin2[~cin2[-1].str.contains('PX')]
        
        if external_select is not None:
            top10f = pd.DataFrame(external_select)
            top10f.columns = ['top10f']
            cin2 = cin2.merge(top10f, left_on=[-1], right_on=['top10f'], how='inner').drop('top10f',axis=1).merge(top10f, left_on=['0_x'], right_on=['top10f'], how='inner').drop('top10f',axis=1)
        
        #cin2['r.sq_diff'] = cin2['r.sq']-cin2['r.sq_max_y']
        if best_plot:           
            cin2 = cin2[cin2['r.sq_diff']>0]
            cin2 = cin2.sort_values('r.sq_diff',ascending=False).groupby(-1).first().reset_index()
        else:
            cin2 = cin2[cin2['r.sq_diff']>min_r2]
            cin2 = cin2.sort_values('r.sq_diff', ascending=False)
            
        print(cin2.shape)
        
        cin2 = cin2.groupby(-1).head(numrow).reset_index()
             
        cin3 = cin2
#         cin3a = cin2.sort_values(['r.sq_diff'], ascending=False).groupby([-1]).head(1).reset_index(drop=True).sort_values(['r.sq_diff'], ascending=False)
#         cin3b = cin2.sort_values(['r.sq_diff'], ascending=False).groupby(['0_x']).head(1).reset_index(drop=True).sort_values(['r.sq_diff'], ascending=False)

#         cin3b = cin3b.drop(-1,axis=1)
#         # Rename columns: _x to _tmp, _y to _x, and _tmp to _y
#         cin3b.columns = [col.replace('_x', '_tmp') for col in cin3b.columns]
#         cin3b.columns = [col.replace('_y', '_x') for col in cin3b.columns]
#         cin3b.columns = [col.replace('_tmp', '_y') for col in cin3b.columns]
#         # Rename the -1 column to '0_x'
#         cin3b[-1] = cin3b['0_y'].copy()
        # cin3 = pd.concat([cin3a,cin3b]).sort_values(['r.sq_diff'], ascending=False).drop_duplicates().head(numrow).reset_index(drop=True)
        # cin3 = cin3.sort_values('r.sq_diff',ascending=False).reset_index(drop=True)

        
        # Create data table
        cin3_table = cin3[[-1,'0_x','r.sq','r.sq_max_y','r.sq_diff']]
        cin3_table.columns = ['Primary', 'Secondary', 'Multi r.sq', 'Uni r.sq', 'r.sq Diff']
        
        # Function to apply formatting only to float columns
        def format_str(val):
            return "{:.4f}".format(val) if isinstance(val, float) else self.translator.custom_translate_omop_2022_2_outtable(val)

        cin3_table = cin3_table.style.format(format_str).set_table_styles([
                                # Aligning all cells to the left
                                {'selector': 'td', 'props': [('text-align', 'left')]},    
                                {'selector': 'th', 'props': [('text-align', 'left')]},  # Aligning index to the left

                                # Optional: Adding a border under the column headers for consistency
                                {'selector': 'thead th', 'props': [('border-bottom', '1px solid black')]},

                                # Style for the caption
                                {'selector': 'caption', 
                                 'props': [('color', 'black'), 
                                           ('background-color', 'yellow'), 
                                           ('font-size', '16px'),
                                           ('text-align', 'left'),
                                           ('font-weight', 'bold')]}, 
                                ])             
        
        
        # Define sizes
        xlabel_size = 14*font_ratio  # Adjust size as needed
        ylabel_size = 20*font_ratio  # Adjust size as needed

        tick_size = 12*font_ratio  # Adjust size as needed
        legend_fontsize = 10*font_ratio  # Adjust size as needed

        if figsize==None:
            figsize=(28, 5 * numrow // ncol)
        
        fig, ax = plt.subplots(numrow // ncol, ncol, figsize=figsize)
        legend_fontsize = 10*font_ratio
        
        
        for i in cin3.index:
            i1 = int(    np.floor(i / ncol)) 
            i2 = int(i - np.floor(i / ncol)*ncol)

            self.plot_meta2_cont(cin3.loc[i, -1], cin3.loc[i, '0_x'], ext_obj=ext_obj, ax=ax[i1][i2], verbose=False, plottype=plottype, contour=contour, ticksize=tick_size, fontsize=xlabel_size)
            
            # Set x_label and y_label for each subplot
            ax[i1][i2].set_xlabel(self.translator.custom_translate_omop_2022_2_outtable(cin3.loc[i, -1]), fontsize=xlabel_size)

            # Increase tick size
            ax[i1][i2].tick_params(axis='both', which='major', labelsize=tick_size)

        # Set a common y_label for all subplots
        fig.text(0.07, 0.5, 'SHAP', va='center', rotation='vertical', fontsize=ylabel_size)    

        fig.savefig(f"meta_double_cont{suffix}.svg", bbox_inches='tight')      
        
        return fig, ax
        
    def plot_r2_rank(self):
        pltable = self.gen_fit_character()[[0, 'r.sq_spline', 'slope_auc']]
        pltable.columns = [0, 'r.sq_spline_AUC', 'slope_auc']
        pltable = pltable.merge(self.gen_fit_character()[[0, 'r.sq_spline']], on=0, how='outer')
        pltable.columns = [0, 'r.sq_spline_AUC', 'slope_auc', 'r.sq_spline_noAUC']
        # pltable = pltable.merge(self_weightAUC.gen_fit_character()[[0, 'r.sq_spline', 'slope_auc']], on=0, how='outer')
        # pltable.columns = [0, 'r.sq_spline_AUC', 'slope_auc', 'r.sq_spline_noAUC', 'r.sq_spline_weightAUC', 'slope_auc_popweight']
        # pltable = pltable.merge(self_weightAUC2.gen_fit_character()[[0, 'r.sq_spline']], on=0, how='outer')
        # pltable.columns = [0, 'r.sq_spline_AUC', 'slope_auc', 'r.sq_spline_noAUC', 'r.sq_spline_weightAUC', 'r.sq_spline_weightAUC2']
        pltable = pltable.sort_values('r.sq_spline_AUC', ascending=False)
        pltable = pltable[~pltable[0].str.contains('PX:')]

        pltable.to_csv('../r2svg.csv')        
        pltable[0] = pltable[0].apply(self.translator.custom_translate_omop_2022_2_fig2)
        
        #plt.plot(pltable[0], pltable['r.sq_spline_AUC'], '-o',  label='AUC')
        #plt.plot(pltable[0], pltable['slope_auc'],       '-o',  label='AUC_slope')
        plt.plot(pltable[0], pltable['r.sq_spline_noAUC'], '-o',  label='no_AUC_popweight')
        
        #plt.plot(pltable[0], pltable['r.sq_spline_weightAUC'], '-+',  label='AUC_popweight')
        #plt.plot(pltable[0], pltable['slope_auc_popweight'], '-x',  label='AUC_slope_popweight')
        # plt.plot(pltable[0], pltable['r.sq_spline_weightAUC2'], '-+',  label='weighted_AUC2')
        #plt.legend()
        plt.xticks(rotation=90)
        #plt.axhline(y=0)
        plt.ylim([0, 1])
        plt.ylabel('$r^2$')
#        plt.title('Best of fit for top features')
        #plt.show()
        plt.savefig('r2.svg', bbox_inches='tight')
                
    def plot_r2_relative_diff(self):
        pltable = self.gen_fit_character()[[0, 'r.sq_spline', 'slope_auc']]
        pltable.columns = [0, 'r.sq_spline_AUC', 'slope_auc']
        pltable = pltable.merge(self.gen_fit_character()[[0, 'r.sq_spline']], on=0, how='outer')
        pltable.columns = [0, 'r.sq_spline_AUC', 'slope_auc', 'r.sq_spline_noAUC']
        pltable = pltable.merge(self_weightAUC.gen_fit_character()[[0, 'r.sq_spline']], on=0, how='outer')
        pltable.columns = [0, 'r.sq_spline_AUC', 'slope_auc', 'r.sq_spline_noAUC', 'r.sq_spline_weightAUC']
        # pltable = pltable.merge(self_weightAUC2.gen_fit_character()[[0, 'r.sq_spline']], on=0, how='outer')
        # pltable.columns = [0, 'r.sq_spline_AUC', 'slope_auc', 'r.sq_spline_noAUC', 'r.sq_spline_weightAUC', 'r.sq_spline_weightAUC2']
        pltable = pltable.sort_values('r.sq_spline_AUC', ascending=False)
        pltable = pltable[~pltable[0].str.contains('PX:')]

        plt.plot(pltable[0], (pltable['r.sq_spline_AUC']-pltable['r.sq_spline_noAUC'])/pltable['r.sq_spline_AUC'], '-o',  label='no_AUC_popweight')
        plt.plot(pltable[0], (pltable['r.sq_spline_AUC']-pltable['r.sq_spline_weightAUC'])/pltable['r.sq_spline_AUC'], '-x',  label='weight_AUC_popweight')
        # plt.plot(pltable[0], (pltable['r.sq_spline_AUC']-pltable['r.sq_spline_weightAUC2'])/pltable['r.sq_spline_AUC'], '-x',  label='weight_AUC2')
        plt.legend()
        plt.xticks(rotation=90)
        plt.axhline(y=0)
        plt.title('r^2 relative difference')
        plt.show()        

    def plot_para(self, target):
        # target = 'intercept_quadratic'
        pltable = self.gen_fit_character()[[0, target, 'r.sq_spline']]
        pltable.columns = [0, target+'_AUC', 'slope_auc']
        pltable = pltable.merge(self.gen_fit_character()[[0, target+'']], on=0, how='outer')
        pltable.columns = [0, target+'_AUC', 'slope_auc', target+'_noAUC']
        pltable = pltable.merge(self_weightAUC.gen_fit_character()[[0, target+'']], on=0, how='outer')
        pltable.columns = [0, target+'_AUC', 'slope_auc', target+'_noAUC', target+'_weightAUC']
        # pltable = pltable.merge(self_weightAUC2.gen_fit_character()[[0, target+'']], on=0, how='outer')
        # pltable.columns = [0, target+'_AUC', 'slope_auc', target+'_noAUC', target+'_weightAUC', target+'_weightAUC2']    
        pltable = pltable.sort_values('slope_auc', ascending=False)
        pltable = pltable[~pltable[0].str.contains('PX:')]

        plt.plot(pltable[0], pltable[target+'_AUC'], '-o',  label='AUC')
        plt.plot(pltable[0], pltable[target+'_noAUC'], '-x',  label='no_AUC_popweight')
        plt.plot(pltable[0], pltable[target+'_weightAUC'], '-+',  label='AUC_popweight')
        # plt.plot(pltable[0], pltable[target+'_weightAUC2'], '-+',  label='weighted_AUC2')    
        plt.legend()
        plt.xticks(rotation=90)
        plt.axhline(y=0)
        plt.title(target)
        plt.show()        
        
    def get_interaction_stat(self, ext_obj, mode='max_y'):
        gamdata_fitdata = self.gamdata_fitdata['double_interaction'].copy()
        gamdata_fitdata['r.sq'] = [x for y in gamdata_fitdata['r.sq'] for x in y]
        cin = gamdata_fitdata.sort_values('r.sq')[[-1,0,'r.sq']].reset_index(drop=True)
        #ext_obj.combine_interaction_json()
        cno = ext_obj.gen_fit_character()
        cno = cno[[0, 'r.sq_linear', 'r.sq_quadratic', 'r.sq_spline']]
        cno['r.sq_max'] = np.max(cno[['r.sq_linear', 'r.sq_quadratic', 'r.sq_spline']],axis=1)
        cin2 = cin.merge(cno, left_on=0, right_on=0,how='left').merge(cno, left_on=-1, right_on=0,how='left')
        cin2['r.sq_max'] = np.max(cin2[['r.sq_max_x', 'r.sq_max_y']],axis=1)

        cin2o = cin2.copy()
        cin2 = cin2.drop(-1,axis=1)
        # Rename columns: _x to _tmp, _y to _x, and _tmp to _y
        cin2.columns = [col.replace('_x', '_tmp') for col in cin2.columns]
        cin2.columns = [col.replace('_y', '_x') for col in cin2.columns]
        cin2.columns = [col.replace('_tmp', '_y') for col in cin2.columns]
        # Rename the -1 column to '0_x'
        cin2[-1] = cin2['0_y'].copy()
        cin2 = cin2[cin2o.columns]
        cin2 = pd.concat([cin2, cin2o])                                
        if mode == 'max_y':
            cin2['r.sq_diff'] = cin2['r.sq']-cin2['r.sq_max_y']
        elif mode == 'max_x':
            cin2['r.sq_diff'] = cin2['r.sq']-cin2['r.sq_max_x']               
        else:
            cin2['r.sq_diff'] = cin2['r.sq']-cin2['r.sq_max']       
        cin2 = cin2.drop_duplicates()
        return cin2
        
    def gen_fit_character(self):
        def reformat_gamdata_fitdata(colname, position, prefix=None):
            def get_all_col(colname, fittype, position=None, df=None, prefix=None):
                if df is None:
                    df = self.gamdata_fitdata[fittype][colname].reset_index(drop=True)
                    if prefix is None:                
                        df.columns = [0] + [x+'_'+fittype for x in colname[1:]]
                    else:
                        df.columns = [0] + [x+'_'+fittype for x in prefix]                
                    if position is not None:
                        df.iloc[:,-1] = [x[position] if position < len(x) else np.nan for x in df.iloc[:,-1]]
                else:
                    dftmp = self.gamdata_fitdata[fittype][colname].reset_index(drop=True)
                    if prefix is None:                
                        dftmp.columns = [0] + [x+'_'+fittype for x in colname[1:]]
                    else:
                        dftmp.columns = [0] + [x+'_'+fittype for x in prefix]                
                    if position is not None:            
                        dftmp.iloc[:,-1] = [x[position] if position < len(x) else np.nan for x in dftmp.iloc[:,-1]]
                    df = df.merge(dftmp, on=0)
                return df

            df = None
            for fittype in ['linear', 'quadratic', 'spline']:
                df = get_all_col(colname, fittype, position, df=df, prefix=prefix)
            return df

        colnamer2 = [0, 'r.sq']
        colnamep = [0, 'p.coeff']
        df_r2 = reformat_gamdata_fitdata(colnamer2, 0)
        df_p_intercept = reformat_gamdata_fitdata(colnamep, 0, prefix=['intercept'])
        df_p_slope = reformat_gamdata_fitdata(colnamep, 1, prefix=['slope'])
        df_p_curvature = reformat_gamdata_fitdata(colnamep, 2, prefix=['curvature'])
        df_p_auc = reformat_gamdata_fitdata(colnamep, -1, prefix=['auc'])
        
        df_all = df_p_intercept.merge(df_p_slope.iloc[:,[0,1,2]], on=0).merge(df_p_curvature.iloc[:,[0,2]], on=0).merge(df_r2, on=0).merge(df_p_auc, on=0)
        df_all = df_all[df_all.columns[[0,1,4,7,2,5,6,8,3,9,10]]]
        df_all = df_all.rename(columns={"auc_linear":"slope_auc"})
        return df_all
    
    def translate_omop_pre(self):
#        path_concept = "/home/hoyinchan/code/concept_vocab/"
        path_concept ="/blue/yonghui.wu/hoyinchan/concept_vocab/"
        self.concept = pd.read_csv(path_concept+'CONCEPT.csv', sep='\t')
        self.concept['vocabulary_id'] = self.concept['vocabulary_id'].replace('CPT4', 'HCPCS')
        self.concept['concept_code'] = self.concept['concept_code'].astype(str)
#        concept_cpt4 = pd.read_csv(path_concept+'CONCEPT_CPT4.csv', sep='\t')
#        concept = pd.concat([concept, concept_cpt4])

    def translate_omop(self, label):
        def get_vocabulary_id(prefix):
            return {
                'LAB': 'LOINC',
                'RX': 'RxNorm',
                'PXCH': 'HCPCS',
                'PX09': 'ICD9Proc',
                'PX10': 'ICD10PCS'
            }.get(prefix, None)

        try:
            prefix, code, extra = label.split(':')
            try:
                extra = extra.split('(')[0]    
            except:
                pass
        #    extra = extra.replace('.','')
            if prefix == 'PX':
                prefix = prefix+code
            label_omop = self.concept[(self.concept['concept_code'] == extra) & (self.concept['vocabulary_id'] == get_vocabulary_id(prefix))]['concept_name'].iloc[0]
        except:
            label_omop = label
        return label_omop
    
    
    #Depreciated
    def extra_translate_omop(self, label):
        # first_dixt = {x:plotshapsn.translate_omop(x) for x in plotshapsn.shapdf['Feature'].unique()}
        # Then ask ChatGPT to extract "extract the medical procedure, medication or Lab measurement from..."
        label = label.split('(')[0]
        data_dict = {
            'LAB::LG6657-3': 'Creatinine (LAB:LONIC:LG6657-3)',
            'LAB::48642-3': 'Glomerular filtration rate predicted among non-blacks (LAB:LONIC:48642-3)',
            'LAB::LG32857-1': 'Leukocytes (LAB:LONIC:LG32857-1)',
            'SYSTOLIC': 'SYSTOLIC',
            'DIASTOLIC': 'DIASTOLIC',
            'PX:CH:36415': 'Collection of venous blood by venipuncture (PX:CH:36415)',
            'LAB::LG6373-7': 'Chloride (LAB:LONIC:LG6373-7)',
            'LAB::LG13614-9': 'Anion gap (LAB:LONIC:LG13614-9)',
            'LAB::48643-1': 'Glomerular filtration rate predicted among blacks (LAB:LONIC:48643-1)',
            'LAB::LG2807-8': 'Bicarbonate (LAB:LONIC:LG2807-8)',
            'LAB::LG49864-8': 'Calcium (LAB:LONIC:LG49864-8)',
            'LAB::LG49936-4': 'Potassium (LAB:LONIC:LG49936-4)',
            'LAB::LG7967-5': 'Glucose (LAB:LONIC:LG7967-5)',
            'LAB::LG4454-7': 'Carbon dioxide (LAB:LONIC:LG4454-7)',
            'LAB::10466-1': 'Anion gap 3 (LAB:LONIC:10466-1)',
            'MED:ATC:B05XA': 'Electrolyte solutions (MED:ATC:B05XA)',
            'LAB::LG1314-6': 'Urea nitrogen (LAB:LONIC:LG1314-6)',
            'AGE': 'AGE',
            'LAB::18182-6': 'Osmolality (LAB:LONIC:18182-6)',
            'ORIGINAL_BMI': 'BMI',
            'LAB::LG49883-8': 'Glucose (LAB:LONIC:LG49883-8)',
            'LAB::LG32892-8': 'Platelets (LAB:LONIC:LG32892-8)',
            'PX:CH:97530': 'Therapeutic activities (PX:CH:97530)',
            'LAB::LG49949-7': 'Phosphate (LAB:LONIC:LG49949-7)',
            'LAB::713-8': 'Eosinophils (LAB:LONIC:713-8)',
            'LAB::LG11363-5': 'Sodium (LAB:LONIC:LG11363-5)',
            'LAB::736-9': 'Lymphocytes (LAB:LONIC:736-9)',
            'PX:CH:97116': 'Gait training (PX:CH:97116)',
            'LAB::LG50041-9': 'Magnesium (LAB:LONIC:LG50041-9)',
            'LAB::1962-0': 'Deprecated Bicarbonate (LAB:LONIC:1962-0)',
            'PX:CH:A6257': 'Transparent film, sterile (PX:CH:A6257)',
            'LAB::LG6139-2': 'Anion gap 4 (LAB:LONIC:LG6139-2)',
            'LAB::2703-7': 'Oxygen partial pressure (LAB:LONIC:2703-7)',
            'LAB::LG6037-8': 'Base deficit (LAB:LONIC:LG6037-8)',
            'LAB::LG344-8': 'Carbon dioxide partial pressure (LAB:LONIC:LG344-8)',
            'LAB::LG12083-8': 'Urea nitrogen/Creatinine ratio (LAB:LONIC:LG12083-8)',
            'SEX_F': 'SEX_F',
            'PX:CH:36556': 'Insertion of central venous catheter (PX:CH:36556)',
            'PX:CH:J8499': 'Prescription drug, oral, non-chemotherapeutic (PX:CH:J8499)',
            'PX:CH:97535': 'Self-care/home management training (PX:CH:97535)',
            'MED:ATC:N02BE': 'Anilides (MED:ATC:N02BE)',
            'PX:CH:J3490': 'Unclassified drugs (PX:CH:J3490)',
            'LAB::786-4': 'MCHC (LAB:LONIC:786-4)',
            'LAB::26478-8': 'Lymphocytes (LAB:LONIC:26478-8)',
            'MED:ATC:A06AA': 'Softeners, emollients (MED:ATC:A06AA)',
            'MED:ATC:L03AC': 'Interleukins (MED:ATC:L03AC)',
            'LAB::LG50477-5': 'Vancomycin trough (LAB:LONIC:LG50477-5)',
            'LAB::LG32850-6': 'Erythrocytes (LAB:LONIC:LG32850-6)',
            'LAB::5905-5': 'Monocytes (LAB:LONIC:5905-5)',
            'LAB::LG44868-4': 'Hemoglobin (LAB:LONIC:LG44868-4)',
            'LAB::LG32885-2': 'Monocytes (LAB:LONIC:LG32885-2)',
            'LAB::788-0': 'Erythrocyte distribution width (LAB:LONIC:788-0)',
            'PX:CH:J3370': 'Injection, vancomycin hcl (PX:CH:J3370)',
            'LAB::LG5665-7': 'Alkaline phosphatase (LAB:LONIC:LG5665-7)',
            'PX:CH:96374': 'Therapeutic, prophylactic, or diagnostic injection (PX:CH:96374)',
            'MED:ATC:A07AA': 'Antibiotics (MED:ATC:A07AA)',
            'PX:CH:99024': 'Postoperative follow-up visit (PX:CH:99024)',
            'LAB::19023-1': 'Granulocytes (LAB:LONIC:19023-1)',
            'LAB::3173-2': 'aPTT (LAB:LONIC:3173-2)',
            'PX:CH:J2543': 'Injection, piperacillin sodium/tazobactam sodium (PX:CH:J2543)',
            'MED:RX:1719287': 'furosemide injection (MED:RX:1719287)',
            'LAB::30385-9': 'Erythrocyte distribution width (LAB:LONIC:30385-9)',
            'DX:09:584bt6': 'Acute kidney failure (DX:09:584bt6)',
            'LAB::4544-3': 'Hematocrit (LAB:LONIC:4544-3)',
            'LAB::14979-9': 'aPTT (LAB:LONIC:14979-9)',
            'LAB::LG1777-4': 'Protein (LAB:LONIC:LG1777-4)',
            'LAB::LG32863-9': 'Lymphocytes (LAB:LONIC:LG32863-9)',
            'MED:ATC:C03CA': 'Sulfonamides, plain (MED:ATC:C03CA)',
            'LAB::770-8': 'Neutrophils (LAB:LONIC:770-8)',
            'LAB::5902-2': 'Prothrombin time (PT) (LAB:LONIC:5902-2)',
            'MED:ATC:B01AB': 'Heparin group (MED:ATC:B01AB)',
            'LAB::20570-8': 'Hematocrit (LAB:LONIC:20570-8)',
            'LAB::LG32849-8': 'Eosinophils (LAB:LONIC:LG32849-8)',
            'LAB::LG6039-4': 'Lactate (LAB:LONIC:LG6039-4)',
            'LAB::32623-1': 'Platelet mean volume (LAB:LONIC:32623-1)',
            'LAB::LG6033-7': 'Aspartate aminotransferase (LAB:LONIC:LG6033-7)',
            'LAB::21000-5': 'Erythrocyte distribution width (LAB:LONIC:21000-5)',
            'PX:CH:96361': 'Intravenous infusion, hydration (PX:CH:96361)',
            'LAB::LG32846-4': 'Granulocytes (LAB:LONIC:LG32846-4)',
            'LAB::LG32886-0': 'Neutrophils (LAB:LONIC:LG32886-0)',
            'PX:CH:J1940': 'Injection, furosemide (PX:CH:J1940)',
            'LAB::787-2': 'MCV (LAB:LONIC:787-2)',
            'LAB::61151-7': 'Albumin (LAB:LONIC:61151-7)',
            'DX:09:276bt6': 'Disorders of fluid electrolyte and acid-base balance (DX:09:276bt6)',
            'LAB::LG49829-1': 'Albumin (LAB:LONIC:LG49829-1)',
            'PX:CH:94002': 'Ventilation assist and management (PX:CH:94002)',
            'LAB::890-4': 'Blood group antibody screen (LAB:LONIC:890-4)',
            'PX:CH:71010': 'Radiologic examination, chest (PX:CH:71010)',
            'LAB::776-5': 'Platelet mean volume (LAB:LONIC:776-5)',
            'LAB::LG50024-5': 'Creatinine (LAB:LOINC:LG50024-5)',
            'LAB::LG49881-2': 'Glucose (LAB:LOINC:LG49881-2)',
            'LAB::LG5903-2': 'Magnesium (LAB:LOINC:LG5903-2)',
            'LAB::LG6426-3': 'Phosphate (LAB:LOINC:LG6426-3)'            
        }
        return data_dict.get(label, label)

    def fit_table_gen_all(self):
        fit_table = self.gen_fit_character()

        def fit_table_gen(fit_table, outputname):
        #    fit_table = fit_table.drop('slope_auc',axis=1)
            fit_table = fit_table.rename({0:'Feature'},axis=1)
            fit_table['Feature'] = fit_table['Feature'].apply(self.translator.custom_translate_omop_2022_2_fig2)
            fit_table.index = fit_table['Feature']
            fit_table = fit_table.drop('Feature',axis=1)
            fit_table = fit_table.drop(['Invasive Mechanical Ventilation', 'Furosemide Injection'])
            fit_table.index.name=None
            
            ## Picture output
            #fit_table = fit_table.T
            fit_table.style.format("{:.4f}").set_table_styles([
                                    # Aligning all cells to the left
                                    {'selector': 'td', 'props': [('text-align', 'left')]},    
                                    {'selector': 'th', 'props': [('text-align', 'left')]},  # Aligning index to the left

                                    # Optional: Adding a border under the column headers for consistency
                                    {'selector': 'thead th', 'props': [('border-bottom', '1px solid black')]},

                                    # Style for the caption
                                    {'selector': 'caption', 
                                     'props': [('color', 'black'), 
                                               ('background-color', 'yellow'), 
                                               ('font-size', '16px'),
                                               ('text-align', 'left'),
                                               ('font-weight', 'bold')]},                    
                            ])
            split_index = [i.split('_')[0] for i in fit_table.columns]
            fit_table.columns = split_index
            dataframe_image.export(fit_table,outputname+".png", table_conversion="playwright")
 #           return fit_table

            ## docx table output
            # Create a new Document
            doc = Document()
            table = doc.add_table(rows=1, cols=len(df.columns) + 1)

            # Set the style of the table if needed
            # doc.styles['Table Grid'].font.size = Pt(12)

            # Insert the column names
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = ''
            hdr_cells[0].paragraphs[0].runs[0].font.bold = True
            hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align
            for i, column in enumerate(df.columns):
                hdr_cells[i+1].text = column
                hdr_cells[i+1].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align

            # Insert the row data
            for index, row in df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(index)
#                row_cells[0].paragraphs[0].runs[0].font.bold = True
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align
                for i, value in enumerate(row):
                    # Formatting numbers with 6 decimal places and center align
                    row_cells[i+1].text = f"{value:.6f}"
                    row_cells[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.save(outputname+'.docx')        
    
    
        fit_table_gen(fit_table[[0, 'intercept_linear', 'slope_linear', 'r.sq_linear']], 'Linear_single')
        fit_table_gen(fit_table[[0, 'intercept_quadratic', 'slope_quadratic', 'curvature_quadratic', 'r.sq_quadratic']], 'Quadratic_single')
        fit_table_gen(fit_table[[0, 'intercept_spline', 'r.sq_spline']], 'Spline_single')    
    
    def gen_interaction_shapalltmp(self, top=10):
        pltable = self.gen_fit_character()[[0, 'r.sq_spline', 'slope_auc']]
        top10_features = pltable.sort_values('r.sq_spline',ascending=False).head(top)[0]
        path = '/home/hoyinchan/blue/Data/data2022/shapalltmp.parquet'
        dfraw =  pd.read_parquet(path)
        get_columns = ['site_m', 'site_d'] + [x+'_Names' for x in top10_features] + [x+'_vals' for x in top10_features]
        
        dfraw = dfraw[get_columns]
        dfraw.to_parquet(path.replace('shapalltmp', 'shapalltmp_interaction'))    
        return top10_features

    def get_interaction_shapalltmp2(self, top=10):
        pltable = self.gen_fit_character()[[0, 'r.sq_spline', 'slope_auc']]
        return pltable.sort_values('r.sq_spline',ascending=False).head(top)[0].reset_index(drop=True)    

# select_type = {'AGE':'quadratic',
#  'PX:09:96.72':'',
#  'PX:CH:J1940':'',
#  'SYSTOLIC':'quadratic',
#  'DIASTOLIC':'quadratic',
#  'ORIGINAL_BMI':'linear',
#  'LAB::17861-6(mg/dL)':'quadratic',
#  'LAB::1975-2(mg/dL)':'linear',
#  'LAB::2075-0(mmol/L)':'quadratic',
#  'LAB::2160-0(mg/dL)':'spline',
#  'LAB::2345-7(mg/dL)':'spline',
#  'LAB::2777-1(mg/dL)':'linear',
#  'LAB::2823-3(mmol/L)':'quadratic',
#  'LAB::2951-2(mmol/L)':'quadratic',
#  'LAB::3094-0(mg/dL)':'spline',
#  'LAB::33037-3(mmol/L)':'linear',
#  'LAB::4092-3(ug/mL)':'quadratic',
#  'LAB::4544-3(%)':'linear',
#  'LAB::6690-2(10*3/uL)':'linear',
#  'LAB::718-7(g/dL)':'spline',
#  'LAB::777-3(10*3/uL)':'quadratic',
#  'LAB::788-0(%)':'spline',
#  'LAB::2028-9(mmol/L)':'spline'}

    def get_minmax(self, feature):
        return self.shapdf[feature+'_Names'].min(), self.shapdf[feature+'_Names'].max()

    def gen_supp_table1_fit(self, fittype):
        outtable = self.gen_fit_character()
        outtable = outtable[~outtable[0].str.contains('PX:')]
        if fittype == 'linear':
            outtable = outtable[[0, 'intercept_'+fittype, 'slope_'+fittype, 'r.sq_'+fittype]]
            outtable.columns = [0, 'intercept', 'slope', 'r.sq']
            outtable = outtable.round(7)
        elif fittype == 'quadratic':
            outtable = outtable[[0, 'intercept_'+fittype, 'slope_'+fittype, 'curvature_'+fittype, 'r.sq_'+fittype]]
            outtable.columns = [0, 'intercept', 'slope', 'curvature', 'r.sq']
            outtable = outtable.round(7)
        else:
            outtable = outtable[[0, 'intercept_'+fittype, 'r.sq_'+fittype]]
            outtable.columns = [0, 'intercept', 'r.sq']
            outtable = outtable.round(7)            
        
        outtable[0] = [self.translator.custom_translate_omop_2022_2_outtable(x) for x in outtable[0]]
        outtable.index = outtable[0]
        outtable = outtable.drop(0,axis=1)
        
        return outtable    

    def gen_supp_table2_fit(self, ext_obj, external_select, best_plot=True, mode='max_y'):

        cin2 = self.get_interaction_stat(ext_obj, mode=mode)
        cin2 = cin2[~cin2['0_x'].str.contains('PX')]
        cin2 = cin2[~cin2[-1].str.contains('PX')]

        if external_select is not None:
            top10f = pd.DataFrame(external_select)
            top10f.columns = ['top10f']
            cin2 = cin2.merge(top10f, left_on=[-1], right_on=['top10f'], how='inner').drop('top10f',axis=1).merge(top10f, left_on=['0_x'], right_on=['top10f'], how='inner').drop('top10f',axis=1)

        #cin2['r.sq_diff'] = cin2['r.sq']-cin2['r.sq_max_y']
        if best_plot:
        #            Swap x and y and combine
            cin2 = cin2[cin2['r.sq_diff']>0]
            cin2 = cin2.sort_values('r.sq_diff',ascending=False).groupby(-1).first().reset_index()
        else:
            cin2 = cin2[cin2['r.sq_diff']>min_r2]

        cin2 = cin2[[-1, '0_x', 'r.sq', 'r.sq_max_y', 'r.sq_max_x', 'r.sq_diff']]
        cin2.columns = ['Primary', 'Secondary', 'r.sq', 'r.sq_max_y', 'r.sq_max_x', 'r.sq_diff']

        cin2['Primary'] = [self.translator.custom_translate_omop_2022_2_outtable(x) for x in cin2['Primary']]
        cin2['Primary'] = cin2['Primary'].str.split('(').str[0]
        cin2['Secondary'] = [self.translator.custom_translate_omop_2022_2_outtable(x) for x in cin2['Secondary']]
        cin2['Secondary'] = cin2['Secondary'].str.split('(').str[0]
        
        cin2.index=cin2['Primary']
        cin2=cin2.drop('Primary',axis=1)
        
        return cin2      
    
    def r2_histogram_interaction(self, ext_obj, mode='max_y'):
        cin2 = self.get_interaction_stat(ext_obj, mode=mode)
        cin2 = cin2[~cin2['0_x'].str.contains('PX')]
        cin2 = cin2[~cin2[-1].str.contains('PX')]
        #cin2['r.sq_diff'] = cin2['r.sq']-cin2['r.sq_max_y']
        cin2 = cin2[cin2['r.sq_diff']>0]

        cin2.sort_values('r.sq_diff', ascending=False)

        #r2 histogram
        df = cin2
        bins = np.arange(0, 1.01, 0.05)
        df['binned'] = pd.cut(df['r.sq_diff'], bins)

        # Plot the histogram
        plt.figure(figsize=(10, 6))
        df['binned'].value_counts().sort_index().plot(kind='bar', edgecolor='black')
        plt.title('Histogram of Binned Normalized Value Column')
        plt.xlabel('Bins')
        plt.ylabel('Frequency')
        plt.grid(False)
        plt.show()
    
    def get_cat_interaction_feature(self):
        xxx = self.gamdata_fitdata['double_interaction']
        cat_fit_feauture = xxx[[x in list(self.cattarget)for x in xxx[0]]]
        cat_fit_feauture['r.sq'] = [x[0] for x in cat_fit_feauture['r.sq']]   
        cat_fit_feauture['intercept'] = [x[0] for x in cat_fit_feauture['p.coeff']]        
        cat_fit_feauture['interaction_intercept'] = [x[1] for x in cat_fit_feauture['p.coeff']]
        return cat_fit_feauture[[-1,'intercept','interaction_intercept', 'r.sq']].drop_duplicates().sort_index()   
    
    def get_cat_interaction_stat(self):
        xxx = self.get_cat_interaction_feature()
        catf = xxx.index.unique()
        all_out=dict()
        for ff in catf:
            dff = xxx[xxx.index==ff]    
            dff.index = dff[-1]
            dff=dff.drop(-1, axis=1)
            dff.index = [self.translator.custom_translate_omop_2022_2_outtable(x) for x in dff.index]
            all_out[ff] = dff 
        return catf, all_out
    
if __name__ == "__main__":
    plotshapsn = plotmeta(order='single', interaction = 'nointeraction')
    plotshapsn.get_meta_data(filename=plotshapsn.home_data_directory+"gamalltmp_single_noAUC.json")
    plotshapsn.load_raw_data()
    plotshapsn.calculate_local_fit_all()
    plotshapsn.cal_plot_range()
    plotshapsn.gen_interaction_shapalltmp(top=17)
    plotshapsn.combine_interaction_json()
    